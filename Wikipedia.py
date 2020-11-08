from docx import Document
import wikipedia
def Wiki(keyword,lang='en'):
                wikipedia.set_lang(lang)
                data = wikipedia.summary(keyword)
                data2 = wikipedia.page(keyword)
                data2 = data2.content
                doc = Document() #Creating docs in Python
                doc.add_heading(keyword,0)
                doc.add_paragraph(data2)
                doc.save(keyword + '.docx')
                print('Sucessfully created')

                
wikipedia.set_lang("en")
from tkinter import *
from tkinter import ttk
from tkinter import messagebox

GUI = Tk()
GUI.title('Wikipedia')
GUI.geometry('400x300')
FONT1  = ('Angsana New',15)
L = ttk.Label(GUI, text="Enter your topic here:",font=FONT1)
L.pack()
v_search = StringVar()
E1 = ttk.Entry(GUI,textvariable=v_search,font=FONT1,width=40)
E1.pack(pady=10)
def Search():
                keyword = v_search.get()#".get" Grabs information from text box
                try:
                                language = v_radio.get()
                                Wiki (keyword,language)
                                messagebox.showinfo('Succesfully recorded')
                except:
                                messagebox.showwarning('Keyword Error','Please enter your search again')
                
                #print (wikipedia.search(keyword))
                #result = wikipedia.summary(keyword,sentences=1)
                #print(result)
B1 = ttk.Button(GUI,text='Search',command=Search)
B1.pack(ipadx=20 ,ipady=10)

#Select language

F1 = Frame(GUI)
F1.pack()
v_radio = StringVar()
RB1 = ttk.Radiobutton(F1, text='Thai',variable=v_radio,value='th')
RB2 = ttk.Radiobutton(F1, text='English',variable=v_radio,value='en')
RB3 = ttk.Radiobutton(F1, text='Chinese',variable=v_radio,value='zh')
RB1.invoke()
RB1.grid(row=0,column=0)
RB2.grid(row=0,column=1)
RB3.grid(row=0,column=2)
GUI.mainloop()

