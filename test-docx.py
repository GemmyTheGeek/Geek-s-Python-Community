from docx import Document
import wikipedia
def Wiki(keyword,lang='en'):
                wikipedia.set_lang(lang)
                data = wikipedia.summary(keyword)
                data2 = wikipedia.page(keyword)
                data2 = data2.content
                doc = Document() #Creating docs in Python
                doc.add_heading(keyword,0)
                doc.add_paragraph(data2)
                doc.save(keyword + '.docx')
                print('Sucessfully created')
Wiki('similan islands')
